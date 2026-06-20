"""
Word导出模块 - 严格遵守PRD
- 代码字体: Consolas
- 填空下划线: Word制表符(Tab) + 下划线前导符
- 代码跨页: 直接截断，依赖防断页校验
- 卷面元素: 密封线/页眉/页脚/考生信息区/大题说明/计分框
- 无水印
"""
import io
import re
import json
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_paragraph(doc, text, font='宋体', size=10.5, bold=False,
                   align=None, indent_cm=None, color=None,
                   space_before=0, space_after=0, line_spacing=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent_cm is not None:
        pf.first_line_indent = Cm(indent_cm)
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _set_font(run, font, size, bold, color)
    return p


def _add_fill_line(doc, width_chars=20):
    """生成填空下划线 - 使用Tab + 下划线前导符（PRD要求绝对平齐）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    # 添加一个Tab + 设置Tab停止位（用下划线作为前导符）
    tab_pos = Cm(width_chars * 0.4)  # 每字符约0.4cm
    p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DASHES)
    # 此处用下划线前导符
    # Word中 UNDERSCORE 前导符即下划线
    # 重新添加
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DASHES)
    # 实际上需要用下划线样式的tab
    # python-docx的WD_TAB_LEADER没有UNDERSCORE，用UNDERScore替代需通过XML
    # 直接使用 DASHES 会有虚线，PRD要求下划线对齐 - 使用XML自定义
    # 重新清空，用XML插入
    p.paragraph_format.tab_stops.clear_all()
    # XML方式
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    # leader="underscore" 类型值
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(width_chars * 0.4 * 567)))  # 1cm=567twips
    tab.set(qn('w:leader'), 'underscore')
    tabs.append(tab)
    pPr.append(tabs)
    run = p.add_run('\t')
    _set_font(run, '宋体', 10.5)
    return p


def _add_code_block(doc, code, font_size=9, is_answer=False, blanks=None):
    """添加代码块 - Consolas字体 + 灰色底纹 + 防跨页 + 支持填空占位符
    is_answer=True 时显示答案，否则显示填空下划线
    """
    if blanks is None:
        blanks = {}
    lines = code.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        _keep_with_next(p, False)
        _add_shading(p, 'F5F5F5')
        if not line:
            line = ' '
        # 处理填空占位符 {{bN}}
        blank_pattern = re.compile(r'\{\{b(\d+)\}\}')
        cursor = 0
        for m in blank_pattern.finditer(line):
            if m.start() > cursor:
                run = p.add_run(line[cursor:m.start()])
                _set_font(run, 'Consolas', font_size, color=RGBColor(0x33, 0x33, 0x33))
            blank_id = f'b{m.group(1)}'
            if is_answer:
                # 答案模式 - 显示答案文本（红色加粗）
                ans_text = blanks.get(blank_id, '【未填】')
                run = p.add_run(ans_text)
                _set_font(run, 'Consolas', font_size, bold=True, color=RGBColor(0xd0, 0x39, 0x39))
            else:
                # 试卷模式 - 填空下划线（Tab + underscore前导符）
                run = p.add_run('\t')
                _set_font(run, 'Consolas', font_size)
                pPr = p._element.get_or_add_pPr()
                tabs = pPr.find(qn('w:tabs'))
                if tabs is None:
                    tabs = OxmlElement('w:tabs')
                    pPr.append(tabs)
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'left')
                tab.set(qn('w:pos'), '1701')  # 约3cm
                tab.set(qn('w:leader'), 'underscore')
                tabs.append(tab)
            cursor = m.end()
        if cursor < len(line):
            run = p.add_run(line[cursor:])
            _set_font(run, 'Consolas', font_size, color=RGBColor(0x33, 0x33, 0x33))


def _add_shading(paragraph, color_hex):
    """段落底纹"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def _keep_with_next(paragraph, keep=True):
    """段落与下段同页"""
    pPr = paragraph._element.get_or_add_pPr()
    e = pPr.find(qn('w:keepNext'))
    if keep and e is None:
        e = OxmlElement('w:keepNext')
        pPr.append(e)
    elif not keep and e is not None:
        pPr.remove(e)


def _keep_lines_together(paragraph, keep=True):
    """段内不分页"""
    pPr = paragraph._element.get_or_add_pPr()
    e = pPr.find(qn('w:keepLines'))
    if keep and e is None:
        e = OxmlElement('w:keepLines')
        pPr.append(e)


def _add_seal_line(doc):
    """左侧密封线 - 通过文本框模拟（简化版：在页眉左边缘加竖排虚线文字）"""
    section = doc.sections[0]
    # 在页眉中添加
    header = section.header
    if not header.paragraphs or not header.paragraphs[0].text:
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # 添加文本框（垂直方向）
        # 简化方案：直接在左侧加竖排文本
        run = p.add_run('  ＿＿＿＿＿＿＿＿  密    封    线    内    不    要    答    题  ＿＿＿＿＿＿＿＿  ')
        _set_font(run, '宋体', 8, color=RGBColor(0x99, 0x99, 0x99))


def _add_student_info(doc, info_text):
    """考生信息区"""
    if not info_text:
        info_text = '姓名：____________   班级：____________   学号：____________   考号：____________'
    _add_paragraph(doc, info_text, font='宋体', size=10.5, bold=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6)


def _add_score_box(doc, big_no_count):
    """计分框 - 表格形式"""
    table = doc.add_table(rows=2, cols=big_no_count + 1)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = table.rows[0].cells
    cells[0].text = '题号'
    for i in range(big_no_count):
        cells[i + 1].text = f'第{i+1}题'
    cells = table.rows[1].cells
    cells[0].text = '得分'
    for i in range(1, big_no_count + 1):
        cells[i].text = ''
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    _set_font(r, '宋体', 10.5, bold=True)
    # 表格后空行
    _add_paragraph(doc, '', size=8)


def _parse_stem_html(html, doc, is_answer=False, blanks=None):
    """解析题干HTML，转换为Word段落
    支持: <img>, <code>, <pre>, 填空占位符{{b1}}
    is_answer=True时，{{b1}} 显示答案；否则显示填空下划线
    """
    if not html:
        return
    # 准备答案映射
    blank_answers = {}
    if blanks:
        for b in blanks:
            if isinstance(b, dict):
                blank_answers[b['id']] = b.get('answer', '')
    # 简化处理：将HTML按行/段拆分
    img_pattern = re.compile(r'<img[^>]+src="([^"]+)"[^>]*/?>(?:</img>)?')
    code_pattern = re.compile(r'<pre[^>]*>(.*?)</pre>', re.DOTALL)
    # 先按<pre>切分
    parts = []
    last_end = 0
    for m in code_pattern.finditer(html):
        parts.append(('text', html[last_end:m.start()]))
        parts.append(('code', m.group(1)))
        last_end = m.end()
    parts.append(('text', html[last_end:]))

    for ptype, content in parts:
        if ptype == 'code':
            # 代码块 - 答案模式下显示答案，否则显示下划线
            _add_code_block(doc, _html_unescape(content), is_answer=is_answer, blanks=blank_answers)
        else:
            # 处理图片
            cursor = 0
            for m in img_pattern.finditer(content):
                if m.start() > cursor:
                    _render_text_with_blanks(doc, content[cursor:m.start()], is_answer, blank_answers)
                # 图片
                src = m.group(1)
                if src.startswith('/static/'):
                    import os
                    img_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), src.lstrip('/'))
                    if os.path.exists(img_path):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_path, width=Cm(12))
                cursor = m.end()
            if cursor < len(content):
                _render_text_with_blanks(doc, content[cursor:], is_answer, blank_answers)


def _render_text_with_blanks(doc, text, is_answer=False, blanks=None):
    """渲染文本，处理填空占位符 {{b1}}"""
    if not text:
        return
    if blanks is None:
        blanks = {}
    # 去掉其他HTML标签
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<p[^>]*>', '\n', text)
    text = re.sub(r'</p>', '', text)
    text = re.sub(r'<strong>(.*?)</strong>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<b>(.*?)</b>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<em>(.*?)</em>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<i>(.*?)</i>', r'\1', text, flags=re.DOTALL)
    # 处理行内代码
    inline_codes = []
    def _save_code(m):
        inline_codes.append(m.group(1))
        return f'\x00CODE{len(inline_codes)-1}\x00'
    text = re.sub(r'<code>(.*?)</code>', _save_code, text, flags=re.DOTALL)
    # 去除剩余标签
    text = re.sub(r'<[^>]+>', '', text)
    text = _html_unescape(text)
    # 按行处理
    lines = text.split('\n')
    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.5
        _keep_with_next(p)
        _keep_lines_together(p)
        # 处理填空占位符
        blank_pattern = re.compile(r'\{\{b(\d+)\}\}')
        cursor = 0
        for m in blank_pattern.finditer(line):
            if m.start() > cursor:
                # 行内代码还原
                seg = line[cursor:m.start()]
                def _restore_code(mm):
                    idx = int(mm.group(1))
                    return inline_codes[idx] if idx < len(inline_codes) else ''
                seg = re.sub(r'\x00CODE(\d+)\x00', _restore_code, seg)
                run = p.add_run(seg)
                _set_font(run, '宋体', 10.5)
            blank_id = f'b{m.group(1)}'
            if is_answer:
                ans_text = blanks.get(blank_id, '【未填】')
                run = p.add_run(f'【答案】{ans_text}')
                _set_font(run, '宋体', 10.5, bold=True, color=RGBColor(0xd0, 0x39, 0x39))
            else:
                # 试卷模式：填空下划线
                run = p.add_run('\t')
                _set_font(run, '宋体', 10.5)
                pPr = p._element.get_or_add_pPr()
                tabs = pPr.find(qn('w:tabs'))
                if tabs is None:
                    tabs = OxmlElement('w:tabs')
                    pPr.append(tabs)
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'left')
                tab.set(qn('w:pos'), '2268')  # 约4cm
                tab.set(qn('w:leader'), 'underscore')
                tabs.append(tab)
            cursor = m.end()
        if cursor < len(line):
            text_part = line[cursor:]
            def _restore_code2(mm):
                idx = int(mm.group(1))
                return inline_codes[idx] if idx < len(inline_codes) else ''
            text_part = re.sub(r'\x00CODE(\d+)\x00', _restore_code2, text_part)
            run = p.add_run(text_part)
            _set_font(run, '宋体', 10.5)


def _html_unescape(s):
    import html as html_mod
    return html_mod.unescape(s)


def export_paper_to_docx(paper, paper_questions, include_answer=False):
    """导出试卷为Word - PRD要求：Consolas代码字体、Tab+下划线填空、防跨页、无水印"""
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 页眉
    if paper.get('header_text'):
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run(paper['header_text'])
        _set_font(run, '宋体', 9, color=RGBColor(0x66, 0x66, 0x66))
    # 页脚
    if paper.get('footer_text'):
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(paper['footer_text'])
        _set_font(run, '宋体', 9, color=RGBColor(0x66, 0x66, 0x66))
    else:
        # 默认页码
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run('— ')
        _set_font(run, '宋体', 9)
        # PAGE字段
        from docx.oxml import OxmlElement
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        run._element.addnext(fld)
        run2 = footer_p.add_run(' —')
        _set_font(run2, '宋体', 9)

    # 密封线
    if paper.get('seal_line', 1):
        _add_seal_line(doc)

    # 标题
    title = paper['title']
    if include_answer:
        title += '（参考答案）'
    _add_paragraph(doc, title, font='黑体', size=18, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    # 副标题
    mode_text = '学考卷' if paper.get('mode') == 'xuekao' else '选考卷'
    _add_paragraph(doc, f'（满分{int(paper["total_score"])}分  {mode_text}）',
                   font='宋体', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    # 考生信息区
    _add_student_info(doc, paper.get('student_info'))
    # 分隔线
    _add_paragraph(doc, '─' * 50, font='宋体', size=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xcc, 0xcc, 0xcc),
                   space_after=6)

    # 计分框
    big_nos = sorted(set(pq['big_no'] for pq in paper_questions if pq['big_no']))
    if big_nos:
        _add_score_box(doc, max(big_nos))

    # 题目
    current_big_no = None
    current_section = None
    for pq in paper_questions:
        # 大题分隔
        if pq['big_no'] != current_big_no:
            current_big_no = pq['big_no']
            # 大题标题
            section_title = pq.get('section_title') or f'第{current_big_no}题'
            _add_paragraph(doc, section_title, font='黑体', size=12, bold=True,
                           space_before=10, space_after=4)
            _keep_with_next(doc.paragraphs[-1])
            # 大题说明
            if pq.get('section_desc'):
                _add_paragraph(doc, pq['section_desc'], font='宋体', size=10.5,
                               space_after=4)
                _keep_with_next(doc.paragraphs[-1])
        # 小题号
        if pq.get('custom_no'):
            q_label = pq['custom_no']
        elif pq['small_no']:
            q_label = f"({pq['small_no']})"
        else:
            q_label = ''
        # 题干
        q_stem = pq.get('q_stem', '')
        # 题号 + 题干
        stem_text = f"{q_label} " if q_label else ''
        # 题型不同处理
        q_type = pq.get('q_type')
        # 单选题需渲染选项
        if q_type == 'single_choice':
            _parse_stem_html(stem_text + q_stem, doc, is_answer=include_answer)
            # 选项
            try:
                options = json.loads(pq.get('q_options') or '{}')
            except:
                options = {}
            opt_lines = []
            for k in ['A', 'B', 'C', 'D']:
                if k in options:
                    opt_lines.append(f"{k}. {options[k]}")
            # 选项排版：每两个一行
            for i in range(0, len(opt_lines), 2):
                line = '    '.join(opt_lines[i:i+2])
                _add_paragraph(doc, line, font='宋体', size=10.5,
                              indent_cm=0.5, space_after=2)
                _keep_with_next(doc.paragraphs[-1])
            if include_answer:
                ans = pq.get('q_answer', '')
                _add_paragraph(doc, f'【答案】{ans}', font='宋体', size=10.5, bold=True,
                              color=RGBColor(0xd0, 0x39, 0x39), indent_cm=0.5, space_after=4)
        else:
            # 非选择题
            try:
                blanks_list = json.loads(pq.get('q_blanks') or '[]')
            except:
                blanks_list = []
            # 答案模式：将blanks映射传给解析函数，{{b1}} 直接显示答案
            _parse_stem_html(stem_text + q_stem, doc, is_answer=include_answer, blanks=blanks_list)
            if include_answer:
                # 如果没有填空，直接显示answer字段
                if not blanks_list:
                    ans = pq.get('q_answer', '')
                    if ans:
                        _add_paragraph(doc, f'【答案】{ans}',
                                       font='宋体', size=10.5, bold=True,
                                       color=RGBColor(0xd0, 0x39, 0x39),
                                       indent_cm=0.5, space_after=4)
                # 解析
                if pq.get('q_analysis'):
                    _add_paragraph(doc, '【解析】', font='宋体', size=10, bold=True,
                                  color=RGBColor(0x33, 0x99, 0x33), indent_cm=0.5, space_after=2)
                    _parse_stem_html(pq['q_analysis'], doc, is_answer=True)
        # 分值标注
        _add_paragraph(doc, f'（{pq.get("score", 0)}分）', font='宋体', size=9,
                      color=RGBColor(0x99, 0x99, 0x99),
                      align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_answer_sheet(paper, pages='single'):
    """答题卡 - 单页/双页模板"""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 密封线
    if paper.get('seal_line', 1):
        _add_seal_line(doc)

    # 标题
    _add_paragraph(doc, f"{paper['title']} 答题卡", font='黑体', size=16, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=6)
    # 考生信息
    _add_student_info(doc, '姓名：____________   班级：____________   学号：____________   考号：____________')
    _add_paragraph(doc, '─' * 50, font='宋体', size=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xcc, 0xcc, 0xcc),
                   space_after=6)

    # 单选题填涂区
    # 取试卷中所有单选题
    from .models import query_all
    pqs = query_all("""SELECT pq.*, q.type AS q_type FROM paper_questions pq
                       JOIN questions q ON q.id=pq.question_id
                       WHERE pq.paper_id=? ORDER BY pq.sort_order""", (paper['id'],))
    single_qs = [pq for pq in pqs if pq['q_type'] == 'single_choice']
    if single_qs:
        _add_paragraph(doc, '一、单选题（请用2B铅笔填涂）', font='黑体', size=12, bold=True,
                       space_before=8, space_after=4)
        # 4列布局
        cols = 4
        rows = (len(single_qs) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for idx, pq in enumerate(single_qs):
            r = idx // cols
            c = idx % cols
            cell = table.rows[r].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            no_text = pq.get('custom_no') or str(pq['small_no'])
            run = p.add_run(f'{no_text}. ')
            _set_font(run, '宋体', 10.5, bold=True)
            for letter in ['A', 'B', 'C', 'D']:
                run = p.add_run(f' [{letter}]')
                _set_font(run, '宋体', 10.5)

    _add_paragraph(doc, '', size=10)

    # 非选择题作答区
    non_single = [pq for pq in pqs if pq['q_type'] != 'single_choice']
    if non_single:
        _add_paragraph(doc, '二、非选择题作答区', font='黑体', size=12, bold=True,
                       space_before=8, space_after=4)
        # 按大题分组
        big_groups = {}
        for pq in non_single:
            big_groups.setdefault(pq['big_no'], []).append(pq)
        for big_no in sorted(big_groups.keys()):
            items = big_groups[big_no]
            section_title = items[0].get('section_title') or f'第{big_no}题'
            _add_paragraph(doc, section_title, font='黑体', size=11, bold=True,
                          space_before=6, space_after=4)
            for pq in items:
                no = pq.get('custom_no') or f"({pq['small_no']})"
                _add_paragraph(doc, f'{no}.', font='宋体', size=10.5, bold=True,
                              space_before=4, space_after=2)
                # 作答空行（下划线）
                for _ in range(6):
                    _add_fill_line(doc, width_chars=45)

    # 如果是双页，分页后继续空白作答区
    if pages == 'double':
        doc.add_page_break()
        _add_paragraph(doc, '续答区', font='黑体', size=12, bold=True,
                       space_before=8, space_after=4)
        for _ in range(20):
            _add_fill_line(doc, width_chars=45)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
